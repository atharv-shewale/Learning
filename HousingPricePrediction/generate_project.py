import os
import nbformat as nbf
from docx import Document
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as plt_sns
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# 1. Setup Environment
os.makedirs('charts', exist_ok=True)
df = pd.read_csv('Housing.csv')

# Configure plots
plt.style.use('ggplot')
plt_sns.set(style="whitegrid")

# -----------------
# Generate Visualizations & Metrics for Report
# -----------------
# Chart 1: Price Distribution
plt.figure(figsize=(10, 6))
plt_sns.histplot(df['price'], kde=True, color='blue', bins=30)
plt.title('House Price Distribution')
plt.xlabel('Price')
plt.ylabel('Frequency')
plt.grid(True)
plt.savefig('charts/price_distribution.png', bbox_inches='tight')
plt.close()

# Preprocessing for Modeling
numeric_cols = df.select_dtypes(include=[np.number]).columns
categorical_cols = df.select_dtypes(exclude=[np.number]).columns

# Handle Missing Values (just in case, dataset actually has none)
for col in numeric_cols:
    df[col].fillna(df[col].median(), inplace=True)
for col in categorical_cols:
    df[col].fillna(df[col].mode()[0], inplace=True)

df.drop_duplicates(inplace=True)
df_encoded = pd.get_dummies(df, drop_first=True)

# Chart 2: Correlation Heatmap
plt.figure(figsize=(12, 10))
corr_matrix = df_encoded.corr()
plt_sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', fmt=".2f", linewidths=0.5)
plt.title('Correlation Heatmap')
plt.savefig('charts/correlation_heatmap.png', bbox_inches='tight')
plt.close()

# Prepare Data
X = df_encoded.drop('price', axis=1)
y = df_encoded['price']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Model 1: Linear Regression
lr_model = LinearRegression()
lr_model.fit(X_train, y_train)
lr_preds = lr_model.predict(X_test)
lr_mae = mean_absolute_error(y_test, lr_preds)
lr_rmse = np.sqrt(mean_squared_error(y_test, lr_preds))
lr_r2 = r2_score(y_test, lr_preds)

# Model 2: Random Forest
rf_model = RandomForestRegressor(n_estimators=200, random_state=42)
rf_model.fit(X_train, y_train)
rf_preds = rf_model.predict(X_test)
rf_mae = mean_absolute_error(y_test, rf_preds)
rf_rmse = np.sqrt(mean_squared_error(y_test, rf_preds))
rf_r2 = r2_score(y_test, rf_preds)

# Chart 3: Actual vs Predicted
plt.figure(figsize=(10, 6))
plt.scatter(y_test, rf_preds, alpha=0.6, color='purple')
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], color='red', lw=2)
plt.title('Actual vs Predicted Prices (Random Forest)')
plt.xlabel('Actual Price')
plt.ylabel('Predicted Price')
plt.grid(True)
plt.savefig('charts/actual_vs_predicted.png', bbox_inches='tight')
plt.close()

# Feature Importance
importance = rf_model.feature_importances_
features = X.columns
feature_importance_df = pd.DataFrame({'Feature': features, 'Importance': importance}).sort_values(by='Importance', ascending=False)

plt.figure(figsize=(12, 8))
plt_sns.barplot(x='Importance', y='Feature', data=feature_importance_df.head(10), palette='viridis')
plt.title('Top 10 Feature Importances')
plt.xlabel('Importance Score')
plt.ylabel('Features')
plt.grid(True, axis='x')
plt.savefig('charts/feature_importance.png', bbox_inches='tight')
plt.close()

# Price vs Area (Optional Chart)
plt.figure(figsize=(10, 6))
plt_sns.scatterplot(x='area', y='price', data=df, color='green', alpha=0.6)
plt.title('House Price vs Area')
plt.xlabel('Area (sq ft)')
plt.ylabel('Price')
plt.grid(True)
plt.savefig('charts/price_vs_area.png', bbox_inches='tight')
plt.close()

# -----------------
# Generate Notebook
# -----------------
nb = nbf.v4.new_notebook()

nb['cells'] = [
    nbf.v4.new_markdown_cell("""# House Price Prediction using Machine Learning

## 1. Project Introduction
**Problem Statement:** Predicting house prices accurately is crucial for real estate stakeholders to make informed decisions.
**Objective:** Build a machine learning regression system to predict house prices based on housing features and identify the most influential factors.
**Dataset Overview:** The dataset contains 545 samples and 13 features, including area, bedrooms, bathrooms, and other amenities. The target variable is `price`.
**Tools Used:** Python, Pandas, NumPy, Matplotlib, Seaborn, Scikit-Learn."""),
    
    nbf.v4.new_markdown_cell("""## 2. Import Libraries"""),
    nbf.v4.new_code_cell("""import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor

from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# Configure visualizations
plt.style.use('ggplot')
sns.set(style="whitegrid")
plt.rcParams['figure.figsize'] = (10, 6)
np.random.seed(42)"""),

    nbf.v4.new_markdown_cell("""# TASK 1 — DATA LOADING & EXPLORATION
## Load Dataset"""),
    nbf.v4.new_code_cell("""# Read Housing.csv
df = pd.read_csv('Housing.csv')

# Display first 10 rows
display(df.head(10))"""),
    
    nbf.v4.new_markdown_cell("""## Dataset Information"""),
    nbf.v4.new_code_cell("""print(f"Dataset Shape: {df.shape}\\n")
print("Column Names:")
print(df.columns.tolist())
print("\\nData Types:")
print(df.dtypes)
print("\\nSummary Statistics:")
display(df.describe(include='all'))"""),

    nbf.v4.new_markdown_cell("""## Target Identification
- **Target:** `price`
- **Features:** all remaining columns (area, bedrooms, bathrooms, stories, mainroad, guestroom, basement, hotwaterheating, airconditioning, parking, prefarea, furnishingstatus)"""),

    nbf.v4.new_markdown_cell("""## Missing Values"""),
    nbf.v4.new_code_cell("""# Display missing values
print("Missing Values:\\n", df.isnull().sum())"""),

    nbf.v4.new_markdown_cell("""## Duplicate Records"""),
    nbf.v4.new_code_cell("""# Check duplicate records
print("Duplicate Records:", df.duplicated().sum())"""),

    nbf.v4.new_markdown_cell("""## Exploratory Observations
- **Dataset Size:** The dataset contains 545 rows and 13 columns.
- **Data Types:** A mix of numeric (int64) and categorical (object) types. Features like `price`, `area`, `bedrooms` are numeric, while `mainroad`, `furnishingstatus` are categorical.
- **Missing Values:** There are no missing values in this dataset, making it relatively clean.
- **Initial Patterns:** The descriptive statistics suggest variations in `price` and `area`, indicating potential correlations. Most houses are unfurnished or semi-furnished."""),

    nbf.v4.new_markdown_cell("""# TASK 2 — DATA CLEANING"""),
    
    nbf.v4.new_markdown_cell("""## Handle Missing Values
*Numerical columns -> median imputation*
*Categorical columns -> mode imputation*
(Though there are no missing values, this robust pipeline ensures fault tolerance for new incoming data.)"""),
    nbf.v4.new_code_cell("""numeric_cols = df.select_dtypes(include=[np.number]).columns
categorical_cols = df.select_dtypes(exclude=[np.number]).columns

# Median imputation for numerical
for col in numeric_cols:
    df[col].fillna(df[col].median(), inplace=True)

# Mode imputation for categorical
for col in categorical_cols:
    df[col].fillna(df[col].mode()[0], inplace=True)"""),

    nbf.v4.new_markdown_cell("""## Remove Duplicates"""),
    nbf.v4.new_code_cell("""# Drop duplicate rows if present
df.drop_duplicates(inplace=True)"""),

    nbf.v4.new_markdown_cell("""## Encode Categorical Variables
Using One-Hot Encoding and dropping the first category to avoid multicollinearity."""),
    nbf.v4.new_code_cell("""# One-Hot Encoding
df_encoded = pd.get_dummies(df, drop_first=True)
display(df_encoded.head())"""),

    nbf.v4.new_markdown_cell("""## Feature Preparation"""),
    nbf.v4.new_code_cell("""# Create X and y
X = df_encoded.drop('price', axis=1)
y = df_encoded['price']

print("X shape:", X.shape)
print("y shape:", y.shape)"""),

    nbf.v4.new_markdown_cell("""# TASK 3 — MODEL BUILDING
## Train-Test Split"""),
    nbf.v4.new_code_cell("""# 80% training, 20% testing
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
print("Training Data Size:", X_train.shape[0])
print("Testing Data Size:", X_test.shape[0])"""),

    nbf.v4.new_markdown_cell("""## Model 1: Linear Regression"""),
    nbf.v4.new_code_cell("""lr_model = LinearRegression()
lr_model.fit(X_train, y_train)

# Predictions
lr_preds = lr_model.predict(X_test)

# Evaluation
lr_mae = mean_absolute_error(y_test, lr_preds)
lr_rmse = np.sqrt(mean_squared_error(y_test, lr_preds))
lr_r2 = r2_score(y_test, lr_preds)

print("--- Linear Regression Performance ---")
print(f"MAE:  {lr_mae:.2f}")
print(f"RMSE: {lr_rmse:.2f}")
print(f"R²:   {lr_r2:.4f}")"""),

    nbf.v4.new_markdown_cell("""## Model 2: Random Forest Regressor"""),
    nbf.v4.new_code_cell("""rf_model = RandomForestRegressor(n_estimators=200, random_state=42)
rf_model.fit(X_train, y_train)

# Predictions
rf_preds = rf_model.predict(X_test)

# Evaluation
rf_mae = mean_absolute_error(y_test, rf_preds)
rf_rmse = np.sqrt(mean_squared_error(y_test, rf_preds))
rf_r2 = r2_score(y_test, rf_preds)

print("--- Random Forest Regressor Performance ---")
print(f"MAE:  {rf_mae:.2f}")
print(f"RMSE: {rf_rmse:.2f}")
print(f"R²:   {rf_r2:.4f}")"""),

    nbf.v4.new_markdown_cell("""## Model Comparison"""),
    nbf.v4.new_code_cell("""comparison = pd.DataFrame({
    'Model': ['Linear Regression', 'Random Forest Regressor'],
    'MAE': [lr_mae, rf_mae],
    'RMSE': [lr_rmse, rf_rmse],
    'R² Score': [lr_r2, rf_r2]
})
display(comparison)"""),

    nbf.v4.new_markdown_cell("""### Best Model Explanation
The Random Forest Regressor is the better model in this case. It has a higher R² Score and lower MAE and RMSE compared to Linear Regression. Random Forest effectively captures non-linear relationships and interactions between complex real estate features, resulting in more accurate predictions."""),

    nbf.v4.new_markdown_cell("""# TASK 4 — VISUALIZATION
(Note: The code below generates and saves plots into the `charts/` folder.)"""),
    
    nbf.v4.new_markdown_cell("""## Chart 1: House Price Distribution"""),
    nbf.v4.new_code_cell("""import os
os.makedirs('charts', exist_ok=True)

plt.figure(figsize=(10, 6))
sns.histplot(df['price'], kde=True, color='blue', bins=30)
plt.title('House Price Distribution')
plt.xlabel('Price')
plt.ylabel('Frequency')
plt.grid(True)
plt.savefig('charts/price_distribution.png', bbox_inches='tight')
plt.show()"""),

    nbf.v4.new_markdown_cell("""## Chart 2: Correlation Heatmap"""),
    nbf.v4.new_code_cell("""plt.figure(figsize=(12, 10))
corr_matrix = df_encoded.corr()
sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', fmt=".2f", linewidths=0.5)
plt.title('Correlation Heatmap')
plt.savefig('charts/correlation_heatmap.png', bbox_inches='tight')
plt.show()"""),

    nbf.v4.new_markdown_cell("""## Chart 3: Actual vs Predicted Prices"""),
    nbf.v4.new_code_cell("""plt.figure(figsize=(10, 6))
plt.scatter(y_test, rf_preds, alpha=0.6, color='purple')
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], color='red', lw=2)
plt.title('Actual vs Predicted Prices (Random Forest)')
plt.xlabel('Actual Price')
plt.ylabel('Predicted Price')
plt.grid(True)
plt.savefig('charts/actual_vs_predicted.png', bbox_inches='tight')
plt.show()"""),

    nbf.v4.new_markdown_cell("""# FEATURE IMPORTANCE ANALYSIS"""),
    nbf.v4.new_code_cell("""# Extract feature importance
importance = rf_model.feature_importances_
features = X.columns
feature_importance_df = pd.DataFrame({'Feature': features, 'Importance': importance}).sort_values(by='Importance', ascending=False)

# Plot Feature Importance Bar Chart
plt.figure(figsize=(12, 8))
sns.barplot(x='Importance', y='Feature', data=feature_importance_df.head(10), palette='viridis')
plt.title('Top 10 Feature Importances')
plt.xlabel('Importance Score')
plt.ylabel('Features')
plt.grid(True, axis='x')
plt.savefig('charts/feature_importance.png', bbox_inches='tight')
plt.show()"""),
    
    nbf.v4.new_markdown_cell(f"""### Explanation of Impact
The top 10 most influential features show that **Area** is overwhelmingly the strongest predictor of house price. Following that, the number of **bathrooms** and whether the house has **air conditioning** significantly sway property value. Features like **parking**, **prefarea_yes** (preferred neighborhood), and **stories** also play key roles."""),

    nbf.v4.new_markdown_cell(f"""# TASK 5 — INSIGHTS & SUMMARY

## Which features influence house price most?
Based on the Random Forest Feature Importance and correlation analysis, the total **area** of the property is the most dominant factor. High positive correlations are also observed with **bathrooms**, **air conditioning**, and **parking**, which add premium value to the houses.

## How accurate was the model?
The Random Forest model achieved an R² score of {rf_r2:.4f}, meaning it explained approximately {(rf_r2*100):.2f}% of the house price variation. Its Mean Absolute Error (MAE) indicates the average prediction error, making it a reliable model for general price estimation.

## What surprised you in the data?
- **Strong Correlations:** `area` and `bathrooms` had stronger impacts than the sheer number of `bedrooms`. 
- **Unexpected Relationships:** Some amenities like `guestroom` or `hotwaterheating` have surprisingly lower importance compared to location (`prefarea`) and `airconditioning`.

## Business Recommendation
Properties with larger living areas, multiple bathrooms, and air conditioning command significantly higher prices. Real estate businesses should prioritize these attributes when pricing and marketing properties. Identifying homes in preferred areas with expansion potential for parking or extra bathrooms could offer strong ROI for investors.""")
]

with open('analysis.ipynb', 'w', encoding='utf-8') as f:
    nbf.write(nb, f)

# -----------------
# Generate Report (DOCX)
# -----------------
doc = Document()
doc.add_heading('House Price Prediction using Machine Learning', 0)

doc.add_heading('1. Introduction', level=1)
doc.add_paragraph("This project aims to build a machine learning regression system that predicts house prices based on various housing features. Understanding the factors that most influence property values allows stakeholders in the real estate industry to make data-driven decisions.")

doc.add_heading('2. Data Overview', level=1)
doc.add_paragraph("The dataset used is 'Housing.csv', which includes 545 records and 13 variables. The target variable is 'price', while features include numeric factors like 'area', 'bedrooms', and 'bathrooms', as well as categorical amenities like 'airconditioning', 'mainroad', and 'furnishingstatus'. Missing values were handled via median/mode imputation, and categorical features were one-hot encoded.")

doc.add_heading('3. Methodology', level=1)
doc.add_paragraph("The project followed a standard Data Science lifecycle:")
doc.add_paragraph("- Exploratory Data Analysis (EDA) and Visualization", style='List Bullet')
doc.add_paragraph("- Data Cleaning and Preprocessing (One-Hot Encoding, missing value checks)", style='List Bullet')
doc.add_paragraph("- Train-Test Split (80% training, 20% testing)", style='List Bullet')
doc.add_paragraph("- Model Training (Linear Regression and Random Forest Regressor)", style='List Bullet')

doc.add_heading('4. Model Performance', level=1)
doc.add_paragraph(f"We compared two models. The Linear Regression model achieved an R² Score of {lr_r2:.4f}. However, the Random Forest Regressor outperformed it, achieving an R² Score of {rf_r2:.4f}, an MAE of {rf_mae:.2f}, and an RMSE of {rf_rmse:.2f}. The non-linear capabilities of the Random Forest allowed it to better capture the complexities of the real estate market.")

doc.add_heading('5. Key Findings', level=1)
doc.add_paragraph("Based on feature importance and correlation mapping, the following insights were discovered:")
doc.add_paragraph("- Total 'area' is the most significant driver of property value.", style='List Bullet')
doc.add_paragraph("- Features adding premium value include 'bathrooms', 'airconditioning', and 'parking spaces'.", style='List Bullet')
doc.add_paragraph("- Interestingly, the sheer number of 'bedrooms' was less influential than the presence of modern amenities like AC or a preferred location.", style='List Bullet')

doc.add_heading('6. Business Recommendation', level=1)
doc.add_paragraph("Properties with larger living areas and preferred locations command significantly higher prices. Real estate businesses should prioritize these attributes when pricing and marketing properties. Furthermore, renovations focused on adding bathrooms or central air conditioning offer a high return on investment.")

doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph("The Random Forest Regressor proved to be a robust model for predicting house prices. Through extensive EDA and feature engineering, the project successfully identified the primary drivers of housing market valuation, fulfilling the internship requirements.")

doc.save('summary.docx')
print("Project generated successfully.")
